from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from .database import Inventory
from .report_consulting import consulting_report_summary

NAVY = "0F2D4D"
GREEN = "2E7D5B"
LIGHT = "EAF3EF"
PALE_BLUE = "EAF4F7"
PALE_YELLOW = "FFF4D6"
PALE_RED = "FDE7E5"
WHITE = "FFFFFF"
TEXT = "17232B"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: object, *, bold: bool = False, color: str = TEXT, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _table(document: Document, rows: list[list[object]], widths: list[float] | None = None):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            _set_cell_text(cell, value, bold=r == 0, color=WHITE if r == 0 else TEXT, size=8.2)
            if r == 0:
                _shade(cell, NAVY)
            elif r % 2 == 0:
                _shade(cell, "F4F7F5")
            if widths and c < len(widths):
                cell.width = Cm(widths[c])
    table.autofit = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def _bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.add_run(text)


def _editable_note(document: Document, title: str, prompt: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, PALE_YELLOW)
    p = cell.paragraphs[0]
    r = p.add_run(f"CAMPO EDITABLE - {title}\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(prompt)
    r2.italic = True
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    document.add_paragraph()


def _configure(document: Document, inventory: Inventory) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 12, GREEN), ("Heading 3", 10, NAVY)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
        styles[name].font.bold = True
    for section in document.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"Calcula tu Huella 1.0.0 | {inventory.name} | Borrador editable sujeto a revisión")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(90, 110, 120)


def generate_editable_consulting_docx(session: Session, inventory: Inventory, output: Path) -> None:
    summary = consulting_report_summary(session, inventory)
    analysis = summary["analysis"]
    delivery = summary["delivery"]
    closure = summary["closure"]
    portfolio = summary["portfolio"]

    document = Document()
    _configure(document, inventory)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CALCULA TU HUELLA")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Informe de huella de carbono - borrador editable de consultoría")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = title.add_run(inventory.organization.name)
    rr.bold = True
    rr.font.size = Pt(20)
    rr.font.color.rgb = RGBColor.from_string(NAVY)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{inventory.name}\n{inventory.start_date:%d/%m/%Y} - {inventory.end_date:%d/%m/%Y}")
    document.add_paragraph()
    _table(document, [
        ["Control", "Valor"],
        ["Versión del inventario", inventory.version],
        ["Estado formal", inventory.status],
        ["Nivel de publicación", delivery["publication"]["level"]],
        ["Alistamiento integral", f"{delivery['score']}%"],
        ["Preparación editorial", f"{summary['report_score']}%"],
        ["Advertencia", "La revisión interna no equivale a verificación independiente."],
    ], [5.5, 11.5])
    _editable_note(document, "mensaje de dirección", "Reemplace este texto por una declaración aprobada por la dirección sobre el propósito y uso del inventario.")
    document.add_page_break()

    _heading(document, "Cómo usar este documento", 1)
    document.add_paragraph(
        "Este archivo es un borrador completamente editable. La plataforma propone una narrativa basada en resultados, calidad y trazabilidad; "
        "el equipo técnico debe validar cada afirmación, completar el contexto organizacional y conservar las limitaciones antes de aprobar o publicar."
    )
    _table(document, [["Capítulo", "Estado", "Puntaje", "Acción"]] + [
        [item["name"], item["status"], f"{item['score']}%", item["action"]] for item in summary["chapters"]
    ], [6.0, 2.7, 2.2, 6.4])

    _heading(document, "1. Resumen ejecutivo", 1)
    document.add_paragraph(delivery["narrative"]["headline"])
    document.add_paragraph(delivery["narrative"]["conclusion"])
    document.add_paragraph(f"Decisión sugerida: {delivery['decision']['primary_decision']}")
    _table(document, [
        ["Indicador", "Resultado"],
        ["Emisiones totales", f"{analysis['total']:.2f} tCO2e"],
        ["Confianza", f"{delivery['decision']['confidence_label']} ({delivery['decision']['confidence_score']}%)"],
        ["Calidad del dato", f"{analysis['quality']['score']}%"],
        ["Cobertura documental", f"{analysis['quality']['evidence_coverage']}%"],
        ["Concentración de las 3 fuentes principales", f"{summary['top_three_share']:.1f}%"],
        ["Preparación del portafolio", f"{portfolio['readiness_score']}%"],
    ], [8.0, 9.0])
    _editable_note(document, "conclusión ejecutiva", "Ajuste la conclusión para el público objetivo. No elimine las cautelas de calidad, comparabilidad o publicación.")

    _heading(document, "2. Perfil, propósito y límites", 1)
    document.add_paragraph(
        f"La organización {inventory.organization.name}, perteneciente al sector {inventory.organization.sector}, desarrolló el inventario "
        f"'{inventory.name}' con el propósito de {inventory.objective.lower()}. El periodo comprende del {inventory.start_date:%d/%m/%Y} "
        f"al {inventory.end_date:%d/%m/%Y}."
    )
    _table(document, [
        ["Elemento", "Definición registrada"],
        ["Metodología", inventory.methodology],
        ["Versión metodológica", inventory.methodology_version],
        ["GWP", inventory.gwp_version],
        ["Consolidación", inventory.consolidation_approach],
        ["Umbral de materialidad", f"{inventory.materiality_threshold}%"],
        ["Año base", inventory.base_year],
    ], [5.5, 11.5])
    _editable_note(document, "límites organizacionales y operacionales", "Explique sedes, operaciones incluidas, exclusiones, cambios de límites y justificaciones materiales.")

    _heading(document, "3. Resultados y materialidad", 1)
    scope_rows = [["Alcance", "tCO2e", "Participación"]]
    for scope in (1, 2, 3):
        value = float(analysis["scopes"].get(scope, 0))
        share = value / analysis["total"] * 100 if analysis["total"] else 0
        scope_rows.append([f"Alcance {scope}", f"{value:.2f}", f"{share:.1f}%"])
    _table(document, scope_rows, [5.0, 5.0, 5.0])
    source_rows = [["Fuente", "Alcance", "Sede", "tCO2e", "%"]]
    for item in analysis["sources_summary"][:10]:
        source_rows.append([item["name"], item["scope"], item["facility"], f"{item['emissions']:.2f}", f"{item['share']:.1f}%"])
    _table(document, source_rows, [5.5, 1.8, 4.0, 2.8, 2.0])

    _heading(document, "4. Comparación e indicadores de intensidad", 1)
    comparison = summary["comparison"]
    if comparison["available"]:
        document.add_paragraph(
            f"Frente a {comparison['previous_year']}, la huella absoluta {comparison['absolute_direction'].lower()} "
            f"{abs(comparison['absolute_change'] or 0):.1f}%. {comparison['warning']}"
        )
    else:
        document.add_paragraph(comparison["warning"])
    intensity_rows = [["Indicador", "Actual", "Anterior", "Variación", "Lectura"]]
    for item in summary["intensities"]:
        intensity_rows.append([
            item["name"],
            "N/D" if item["value"] is None else f"{item['value']:.6f} {item['unit']}",
            "N/D" if item["previous_value"] is None else f"{item['previous_value']:.6f}",
            "N/D" if item["change"] is None else f"{item['change']:+.1f}%",
            item["direction"],
        ])
    _table(document, intensity_rows, [3.0, 4.0, 3.0, 2.4, 3.0])
    _editable_note(document, "explicación de variaciones", "Separe cambios por actividad, eficiencia, límites, factores, adquisiciones, cierres o calidad del dato.")

    _heading(document, "5. Calidad, evidencia e incertidumbre", 1)
    _table(document, [
        ["Dimensión", "Resultado", "Lectura"],
        ["Calidad consolidada", f"{analysis['quality']['score']}%", "Ponderación de niveles A-D de los registros."],
        ["Cobertura documental", f"{analysis['quality']['evidence_coverage']}%", "Registros con evidencia vinculada."],
        ["Datos estimados", f"{analysis['quality']['estimated_share']}%", "Participación de registros marcados como estimados."],
        ["Incertidumbre combinada", f"{closure['uncertainty']['combined_percentage']:.2f}%", "Aplica sobre las emisiones cubiertas."],
        ["Cobertura de incertidumbre", f"{closure['uncertainty']['emission_coverage_percentage']:.1f}%", "Porción del inventario bruto con cuantificación."],
    ], [5.0, 3.5, 8.5])

    _heading(document, "6. Hallazgos e interpretación", 1)
    _table(document, [["Prioridad", "Tema", "Hallazgo", "Evidencia", "Implicación"]] + [
        [item["level"], item["topic"], item["finding"], item["evidence"], item["implication"]] for item in summary["findings"]
    ], [2.2, 2.4, 4.2, 4.1, 4.5])

    _heading(document, "7. Recomendaciones y plan de acción", 1)
    _table(document, [["Prioridad", "Tema", "Acción", "Responsable sugerido", "Criterio de aceptación"]] + [
        [item["priority"], item["topic"], item["action"], item["owner"], item["acceptance"]] for item in summary["recommendations"]
    ], [2.2, 2.3, 4.6, 3.8, 4.5])

    _heading(document, "8. Portafolio de reducción", 1)
    document.add_paragraph(
        f"Estado: {portfolio['portfolio_status']}. Cobertura de la meta: {portfolio['coverage_percent']:.1f}%. "
        f"Preparación: {portfolio['readiness_score']}%. Decisión principal: {portfolio['primary_decision']}"
    )
    reduction_rows = [["Clasificación", "Acción", "Reducción esperada", "Preparación", "Responsable", "Estado"]]
    for item in portfolio["actions"][:12]:
        reduction_rows.append([
            item["classification"], item["title"], f"{item['expected_reduction']:.2f} tCO2e/año",
            f"{item['readiness_score']}%", item["owner"], item["status"],
        ])
    if len(reduction_rows) == 1:
        reduction_rows.append(["Pendiente", "Crear portafolio de reducción", "N/D", "0%", "Dirección", "No iniciado"])
    _table(document, reduction_rows, [3.0, 4.8, 3.2, 2.2, 3.3, 2.5])

    _heading(document, "9. Limitaciones y reglas de comunicación", 1)
    for item in summary["limitations"]:
        _bullet(document, f"{item['category']}: {item['detail']}")
    _table(document, [["Afirmación", "Permitida", "Orientación"]] + [
        [item["label"], "Sí" if item["allowed"] else "No", item["guidance"]] for item in summary["claims"]
    ], [4.0, 2.0, 10.5])
    _editable_note(document, "uso previsto", "Indique quién recibirá el documento, para qué decisión y bajo qué control de confidencialidad.")

    _heading(document, "10. Anexo metodológico", 1)
    _table(document, [
        ["Partida", "Valor", "Tratamiento"],
        ["Emisiones brutas", f"{closure['balance']['gross_emissions']:.3f} tCO2e", "Inventario corporativo"],
        ["CO2 biogénico", f"{closure['balance']['biogenic_memo']:.3f} tCO2e", "Partida informativa"],
        ["Remociones", f"{closure['balance']['removals']:.3f} tCO2e", "Separadas del bruto"],
        ["Emisiones evitadas", f"{closure['balance']['avoided_emissions']:.3f} tCO2e", "Fuera del inventario físico"],
        ["Compensaciones", f"{closure['balance']['offsets']:.3f} tCO2e", "Fuera del inventario bruto"],
    ], [5.0, 4.0, 7.5])
    document.add_paragraph(
        "La memoria de cálculo, los factores versionados, los datos de actividad, las evidencias, las decisiones y la pista de auditoría "
        "se conservan en el expediente de la plataforma y deben acompañar cualquier proceso de revisión o verificación."
    )

    document.save(output)
