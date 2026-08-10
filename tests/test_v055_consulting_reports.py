from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import select

from app.database import Base, ENGINE, Inventory, ReportArtifact, SessionLocal, init_db
from app.main import app
from app.report_consulting import consulting_report_summary, portfolio_control_view
from app.report_docx import generate_editable_consulting_docx
from app.reporting import generate_calculation_workbook
from app.services.reports import generate_report
from app.storage import storage


@pytest.fixture(autouse=True)
def fresh_database():
    # La base semilla aislada se restaura desde tests/conftest.py.
    yield


def login(client: TestClient, email: str = "consultor@calculatuhuella.local") -> None:
    response = client.post(
        "/login",
        data={"email": email, "password": "Demo2026!"},
        follow_redirects=False,
    )
    assert response.status_code == 303


def test_v055_summary_explains_comparison_intensity_and_claims():
    with SessionLocal() as session:
        inventory = session.scalar(select(Inventory).where(Inventory.id == 1))
        summary = consulting_report_summary(session, inventory)
        assert summary["version"] == "1.0.0"
        assert 0 <= summary["report_score"] <= 100
        assert len(summary["chapters"]) == 7
        assert len(summary["intensities"]) == 3
        assert summary["findings"]
        assert summary["limitations"]
        control = {item["code"]: item for item in summary["portfolio_control"]}
        assert control["required_reduction"]["value"] == summary["portfolio"]["required_reduction"]
        assert control["expected_reduction"]["value"] == summary["portfolio"]["expected_reduction"]
        assert control["gap"]["value"] == summary["portfolio"]["gap"]
        assert control["coverage_percent"]["value"] == summary["portfolio"]["coverage_percent"]
        assert any(item["label"] == "Inventario verificado" and not item["allowed"] for item in summary["claims"])


def test_v055_portfolio_control_handles_zero_gap_and_overcoverage_without_claiming_compliance():
    zero = portfolio_control_view({
        "required_reduction": 0.0, "expected_reduction": 0.0, "gap": 0.0, "coverage_percent": 0.0,
    })
    assert {item["code"]: item["value"] for item in zero} == {
        "required_reduction": 0.0, "expected_reduction": 0.0, "gap": 0.0, "coverage_percent": 0.0,
    }

    over = portfolio_control_view({
        "required_reduction": 100.0, "expected_reduction": 120.0, "gap": 0.0, "coverage_percent": 120.0,
    })
    coverage = next(item for item in over if item["code"] == "coverage_percent")
    assert coverage["value"] == 120.0
    assert "no implica cumplimiento automático" in coverage["reading"]


def test_v055_workshop_page_and_api_load():
    with TestClient(app) as client:
        login(client)
        page = client.get("/reportes/consultoria")
        assert page.status_code == 200
        assert "Taller de informe" in page.text
        assert "Resultado relacionado con la escala de operación" in page.text
        assert "Generar Word editable" in page.text
        api = client.get("/api/reportes/consultoria")
        assert api.status_code == 200
        payload = api.json()
        assert payload["version"] == "1.0.0"
        assert payload["inventory_id"] == 1
        assert len(payload["chapters"]) == 7


def test_v055_editable_docx_is_valid_and_substantive(tmp_path: Path):
    output = tmp_path / "consulting.docx"
    with SessionLocal() as session:
        inventory = session.scalar(select(Inventory).where(Inventory.id == 1))
        generate_editable_consulting_docx(session, inventory, output)
    assert output.stat().st_size > 10_000
    with ZipFile(output) as archive:
        assert "word/document.xml" in archive.namelist()
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "Informe de huella de carbono" in text
    assert "Resumen ejecutivo" in text
    assert "Gobierno de la entrega, limitaciones y uso" in text
    assert "Declaración técnica y próximos pasos" in text
    for label in ("Reducción requerida", "Reducción esperada", "Brecha de reducción", "Cobertura del portafolio"):
        assert label in table_text
    assert len(document.tables) >= 8


def test_v055_editable_artifact_is_persisted_and_downloadable():
    with SessionLocal() as session:
        inventory = session.scalar(select(Inventory).where(Inventory.id == 1))
        artifact = generate_report(session, inventory, "editable", actor_email="consultor@calculatuhuella.local")
        session.commit()
        artifact_id = artifact.id
        stored_name = artifact.stored_name
        assert artifact.file_name.endswith(".docx")
        assert artifact.report_type == "Informe de consultoría editable"
        assert artifact.file_size > 10_000
        assert storage.exists(stored_name)
        persisted = session.get(ReportArtifact, artifact_id)
        assert persisted is not None
        assert len(persisted.sha256) == 64


def test_v055_workbook_contains_consulting_narrative(tmp_path: Path):
    output = tmp_path / "calculation.xlsx"
    with SessionLocal() as session:
        inventory = session.scalar(select(Inventory).where(Inventory.id == 1))
        generate_calculation_workbook(session, inventory, output)
    workbook = load_workbook(output, read_only=True, data_only=False)
    assert "Narrativa consultoría" in workbook.sheetnames
    values = [cell for row in workbook["Narrativa consultoría"].iter_rows(values_only=True) for cell in row if cell]
    joined = " | ".join(str(value) for value in values)
    assert "CAPÍTULOS Y PREPARACIÓN EDITORIAL" in joined
    assert "HALLAZGOS" in joined
    assert "REGLAS DE COMUNICACIÓN" in joined


def test_v055_reports_page_exposes_five_deliverables():
    with TestClient(app) as client:
        login(client)
        response = client.get("/reportes")
        assert response.status_code == 200
        assert "V0.55 · informe explicable" in response.text
        assert "Informe de consultoría editable" in response.text
        assert response.text.count('name="report_type"') >= 5
